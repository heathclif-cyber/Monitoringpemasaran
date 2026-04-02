import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C = 'Calibri'
FS = 9

def _no_borders(tbl_or_cell, is_table=True):
    sides = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    if is_table:
        elem = tbl_or_cell._tbl
        pr = elem.find(qn('w:tblPr'))
        if pr is None:
            pr = OxmlElement('w:tblPr')
            elem.insert(0, pr)
        tag_str = 'w:tblBorders'
        tag = qn(tag_str)
    else:
        elem = tbl_or_cell._tc
        pr = elem.get_or_add_tcPr()
        tag_str = 'w:tcBdr'
        tag = qn(tag_str)
    old = pr.find(tag)
    if old is not None: pr.remove(old)
    bdr = OxmlElement(tag_str)
    for s in sides:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0'); b.set(qn('w:color'), 'auto')
        bdr.append(b)
    pr.append(bdr)

def _cw(cell, cm):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    old = pr.find(qn('w:tcW'))
    if old is not None: pr.remove(old)
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(cm * 567)))
    w.set(qn('w:type'), 'dxa')
    pr.append(w)

def _valign_top(cell):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), 'top')
    pr.append(v)

def _run(para, text, bold=False, size=FS, underline=False, italic=False):
    r = para.add_run(str(text))
    r.bold = bold; r.underline = underline; r.italic = italic
    r.font.name = C; r.font.size = Pt(size)
    return r

def _para_fmt(p, sa=1):
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(sa)

def _cp(cell):
    p = cell.paragraphs[0]; _para_fmt(p); return p

def _s(v, fb='-'):
    return str(v).strip() if v else fb

def _id_fmt(v, decimals=2):
    if v is None: return '-'
    res = ("{:,." + str(decimals) + "f}").format(v).replace(",", "X").replace(".", ",").replace("X", ".")
    return res

def _rp(v):
    if v is None: return '-'
    return 'Rp' + _id_fmt(v, 2)

def _rp_full(v):
    if not v: return 'Rp0,00'
    return 'Rp' + _id_fmt(v, 2)

MONTHS = ['', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
def _date(d):
    if not d: return '-'
    try: return f"{d.day} {MONTHS[d.month]} {d.year}"
    except: return str(d)

def generate_contract_docx(k) -> io.BytesIO:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2)

    vol = k.volume or 0
    harga = k.harga_satuan or 0
    premi = k.premi or 0
    ppn = k.ppn_persen or 11
    sat = _s(k.satuan, 'Unit')
    nilai = k.nilai_transaksi or 0

    vol_str = _id_fmt(vol, 2) + ' ' + sat if vol else '-'
    harga_str = _rp_full(harga) + ' per ' + sat if harga else '-'
    premi_str = _rp(premi)
    jml_str = _rp_full(nilai) + (f" ({_s(k.terbilang)} Rupiah)" if k.terbilang else "")
    lama = k.lama_pembayaran_hari or 15

    pjl_lines = _s(k.penjual).replace('\\n', '\n').split('\n')
    pjl_nama = pjl_lines[0].strip()
    pjl_addr = pjl_lines[1].strip() if len(pjl_lines) > 1 else 'Jalan Urip Sumoharjo No. 72-76, Kota Makassar'

    pbl_lines = _s(k.pembeli, '[Nama Pembeli]').replace('\\n', '\n').split('\n')
    pbl_nama = pbl_lines[0].strip()
    pbl_addr = pbl_lines[1].strip() if len(pbl_lines) > 1 else _s(k.alamat_pembeli, '')
    if pbl_addr == '-': pbl_addr = ''

    syarat = [ln.strip() for ln in _s(k.syarat_syarat, '').replace('\\n', '\n').split('\n') if ln.strip()]
    if not syarat:
        ambil = k.penyerahan_hari or 15
        syarat = [
            f"Pembayaran dilaksanakan selambat-lambatnya {lama} hari kalender setelah ditandatanganinya Kontrak penjualan",
            "Penyerahan barang dilaksanakan setelah diterima bukti transfer pembayaran",
            f"Pengambilan barang selambat-lambatnya {ambil} hari kalender dari batas akhir tanggal pembayaran",
            "Biaya-biaya yang terkait dengan administrasi Bank menjadi beban pembeli",
        ]

    dasar = _s(k.dasar_ketentuan, "Mengacu kepada tata Cara dan Ketentuan Penjualan Komoditi Perkebunan PT Perkebunan Nusantara III (Persero)")
    lokasi = _s(k.lokasi, 'Makassar')
    tgl = _date(k.tanggal_kontrak)

    met = _s(k.pembayaran_metode, 'Tunai')
    cara = _s(k.pembayaran_cara, 'Transfer')
    bank = _s(k.pembayaran_bank, 'Bank Rakyat Indonesia')
    atas = _s(k.pembayaran_atas_nama, 'PT Perkebunan Nusantara I Regional 8')
    rek = _s(k.pembayaran_rek_no, 'No. 0050-01-005356-30-0')

    def ctr(text, bold=False, ul=False, size=FS, sa=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_fmt(p, sa=sa)
        _run(p, text, bold=bold, underline=ul, size=size)

    ctr('KONTRAK PENJUALAN', bold=True, ul=True, size=11, sa=1)
    ctr(f"Nomor : {_s(k.no_kontrak)}", sa=1)
    ctr(f"Bid Offer Nomor : {_s(k.bid_offer_nomor, '-')}", sa=6)

    CW = [3.8, 0.25, 4.8, 2.8, 0.25, 3.6]
    tbl = doc.add_table(rows=0, cols=6)
    _no_borders(tbl, is_table=True)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    def nr():
        row = tbl.add_row()
        for i, cell in enumerate(row.cells):
            _no_borders(cell, False); _cw(cell, CW[i]); _valign_top(cell); _para_fmt(cell.paragraphs[0])
        return row

    def W(cell, text, bold=False): _run(_cp(cell), text, bold=bold)
    def COLON(cell): _run(_cp(cell), ':')

    def RS(label, value, bv=False):
        r = nr(); W(r.cells[0], label, True); COLON(r.cells[1])
        r.cells[2].merge(r.cells[5]); W(r.cells[2], value, bv)

    def RM(label, lines):
        r = nr(); W(r.cells[0], label, True); COLON(r.cells[1]); r.cells[2].merge(r.cells[5])
        p = _cp(r.cells[2])
        for i, (txt, b) in enumerate(lines):
            if i: p.add_run('\n')
            _run(p, txt, b)

    def RD(l1, v1, l2, v2, bv1=False, bv2=False):
        r = nr(); W(r.cells[0], l1, True); COLON(r.cells[1]); W(r.cells[2], v1, bv1)
        W(r.cells[3], l2, True); COLON(r.cells[4]); W(r.cells[5], v2, bv2)

    RS('Pemilik Komoditas', _s(k.pemilik_komoditas), True)
    RM('Penjual', [(pjl_nama, True), (pjl_addr, False)])
    lines_pb = [(pbl_nama, True)]
    if pbl_addr: lines_pb.append((pbl_addr, False))
    RM('Pembeli', lines_pb)
    RS('No. Referensi', _s(k.no_reff))
    RD('Komoditi', _s(k.komoditi), 'Jenis Komoditi', _s(k.jenis_komoditi))
    RD('Packaging', _s(k.packaging), 'Symbol', _s(k.simbol))
    RS('Deskripsi Produk', _s(k.deskripsi_produk))
    RS('Mutu', _s(k.mutu))
    RS('Produsen', _s(k.kebun_produsen))
    RS('Pelabuhan Muat', _s(k.pelabuhan_muat))
    RS('Volume', vol_str)
    RD('Harga Satuan', harga_str, 'Premi', premi_str)
    ppn_label = f"Tarif Efektif {ppn}%"
    if str(getattr(k, 'is_ppn', 'true')).lower() == 'false':
        ppn_label = "Non-PPN (Bebas PPN)"
    RS('PPN', ppn_label)
    RS('Kondisi Penyerahan', _s(k.kondisi_penyerahan))
    
    # Pembayaran
    r1 = nr()
    W(r1.cells[0], 'Pembayaran', True); COLON(r1.cells[1])
    p2 = _cp(r1.cells[2])
    _run(p2, 'Metode', True); _run(p2, '    :    '); _run(p2, met)
    W(r1.cells[3], 'Cara\nPembayaran', True); COLON(r1.cells[4]); W(r1.cells[5], cara)

    r2 = nr()
    p2b = _cp(r2.cells[2])
    _run(p2b, 'Nama Bank', True); _run(p2b, '   :   '); _run(p2b, bank)
    W(r2.cells[3], 'Jatuh Tempo\nPembayaran', True); COLON(r2.cells[4])
    W(r2.cells[5], f"Maksimal {lama} Hari\nKalender")

    r3 = nr(); r3.cells[2].merge(r3.cells[5]); p3 = _cp(r3.cells[2])
    _run(p3, 'Atas Nama', True); _run(p3, '  :  '); _run(p3, atas)
    r4 = nr(); r4.cells[2].merge(r4.cells[5]); p4 = _cp(r4.cells[2])
    _run(p4, 'Rek No.', True); _run(p4, '    :    '); _run(p4, rek)

    RS('Waktu Penyerahan', _s(k.waktu_penyerahan))
    
    r = nr()
    W(r.cells[0], 'Syarat - Syarat Lain', True); COLON(r.cells[1])
    r.cells[2].merge(r.cells[5]); p = _cp(r.cells[2])
    alpha = 'abcdefgh'
    for i, s in enumerate(syarat):
        clean = s.strip()
        if len(clean) > 2 and clean[0].lower() in alpha and clean[1] in '.':
            clean = clean[2:].strip()
        letter = alpha[i] if i < len(alpha) else str(i+1)
        if i: p.add_run('\n')
        _run(p, f"{letter}.    {clean}")

    RS('Dasar Ketentuan', dasar)
    RS('Jumlah Pembayaran', jml_str)
    RS('Catatan', '-')

    # Signatures (Buyer only)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    sig = doc.add_table(rows=4, cols=1)
    _no_borders(sig, True); sig.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r in sig.rows:
        for c in r.cells: _no_borders(c, False); _para_fmt(c.paragraphs[0])
    
    p_date = sig.rows[0].cells[0].paragraphs[0]
    _run(p_date, f"{lokasi}, {tgl}")
    
    p_l = sig.rows[1].cells[0].paragraphs[0]; _run(p_l, 'Persetujuan Pembeli', True)
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_invoice_docx(invoice) -> io.BytesIO:
    from datetime import timedelta
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Cm(1.5); sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)

    k = invoice.kontrak

    CW = [1.0, 3.0, 1.5, 1.5, 2.0, 1.5, 2.0, 2.0, 0.5, 3.0]
    tbl = doc.add_table(rows=0, cols=10)

    # helper for borders
    elem = tbl._tbl
    pr = elem.find(qn('w:tblPr'))
    if pr is None:
        pr = OxmlElement('w:tblPr')
        elem.insert(0, pr)
    tag_str = 'w:tblBorders'
    bdr = OxmlElement(tag_str)
    for s in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        bdr.append(b)
    pr.append(bdr)

    def nr():
        row = tbl.add_row()
        for i, c in enumerate(row.cells):
            tc = c._tc
            pr = tc.get_or_add_tcPr()
            # remove existing width
            old_w = pr.find(qn('w:tcW'))
            if old_w is not None: pr.remove(old_w)
            w = OxmlElement('w:tcW')
            w.set(qn('w:w'), str(int(CW[i] * 567)))
            w.set(qn('w:type'), 'dxa')
            pr.append(w)
            
            # center vertically by default
            v = pr.find(qn('w:vAlign'))
            if v is None:
                v = OxmlElement('w:vAlign')
                pr.append(v)
            v.set(qn('w:val'), 'top')
            _para_fmt(c.paragraphs[0], 0)
        return row

    def add_rt(c, l, v, bf=True, bsv=False):
        p = c.paragraphs[0]
        _run(p, l, bold=bf)
        if v:
            p.add_run('\n')
            _run(p, v, bold=bsv)

    # 17 rows
    rows = [nr() for _ in range(17)]

    # Row 0 to 3 Left
    c_kepada = rows[0].cells[0]
    c_kepada.merge(rows[3].cells[4])
    p_kepada = c_kepada.paragraphs[0]
    _run(p_kepada, 'Kepada Yth:\n', bold=True)
    pbl_lines = _s(k.pembeli, '[Nama Pembeli]').replace('\\n', '\n').split('\n')
    _run(p_kepada, f"{pbl_lines[0].strip()}\n", bold=True)
    _run(p_kepada, 'Di –\n    Tempat')

    # Row 0 Right
    c_prof = rows[0].cells[5]
    c_prof.merge(rows[0].cells[9])
    p_prof = c_prof.paragraphs[0]
    v_a = c_prof._tc.get_or_add_tcPr().find(qn('w:vAlign'))
    if v_a is not None: v_a.set(qn('w:val'), 'center')
    _run(p_prof, 'Proforma Invoice', bold=True, size=18)

    # Row 1 Right
    c_prof_no = rows[1].cells[5]
    c_prof_no.merge(rows[1].cells[9])
    add_rt(c_prof_no, 'Proforma Invoice No:', invoice.no_invoice)

    # Row 2 Right
    c_kontrak_no = rows[2].cells[5]
    c_kontrak_no.merge(rows[2].cells[9])
    add_rt(c_kontrak_no, 'Kontrak No:', k.no_kontrak)

    # Row 3 Right
    c_ref = rows[3].cells[5]
    c_ref.merge(rows[3].cells[9])
    add_rt(c_ref, 'No. Ref Pembeli:', _s(k.no_reff))

    # Row 4
    c_prod = rows[4].cells[0]
    c_prod.merge(rows[4].cells[4])
    add_rt(c_prod, 'Produsen:', _s(k.kebun_produsen, '-'))

    c_tgl = rows[4].cells[5]
    c_tgl.merge(rows[4].cells[9])
    add_rt(c_tgl, 'Tanggal:', _date(invoice.tanggal_transaksi))

    # Row 5 to 7 Left
    c_desc = rows[5].cells[0]
    c_desc.merge(rows[7].cells[4])
    add_rt(c_desc, 'Deskripsi Produk:', _s(k.komoditi), bf=True, bsv=True)

    # Row 5 Right
    c_jt = rows[5].cells[5]
    c_jt.merge(rows[5].cells[9])
    jt_date = '-'
    if invoice.tanggal_transaksi:
        jt_date = _date(invoice.tanggal_transaksi + timedelta(days=k.lama_pembayaran_hari or 15))
    add_rt(c_jt, 'Tanggal Jatuh Tempo:', jt_date)

    # Row 6 Right
    c_mutu = rows[6].cells[5]
    c_mutu.merge(rows[6].cells[9])
    add_rt(c_mutu, 'Mutu:', _s(k.mutu))

    # Row 7 Right
    c_ppn = rows[7].cells[5]
    c_ppn.merge(rows[7].cells[9])
    ppn_pct = k.ppn_persen or 11
    add_rt(c_ppn, 'PPN:', f"Pajak Pertambahan Nilai (Tarif Efektif {ppn_pct}%)")

    # Row 8
    c_kon = rows[8].cells[0]
    c_kon.merge(rows[8].cells[4])
    add_rt(c_kon, 'Kondisi Penyerahan:', _s(k.kondisi_penyerahan))

    c_pel = rows[8].cells[5]
    c_pel.merge(rows[8].cells[9])
    add_rt(c_pel, 'Pelabuhan Muat:', _s(k.pelabuhan_muat))

    # Row 9 (Info text)
    c_info = rows[9].cells[0]
    c_info.merge(rows[9].cells[9])
    p_info = c_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_info, 'Untuk dan atas nama PT Perkebunan Nusantara I, harap dilakukan pembayaran atas penyerahan:')

    # Row 10 (Header Items)
    headers = ['No.\nReff', 'Jenis\nKomoditi', 'Packaging', 'Chop', 'Symbol/\nKebun', 'Pack\nQty', 'Volume\n(Kg)', 'Harga\n(IDR)', 'Nilai Rupiah']
    for i in range(8):
        c = rows[10].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(c.paragraphs[0], headers[i], bold=True)
        v = c._tc.get_or_add_tcPr().find(qn('w:vAlign'))
        if v is not None: v.set(qn('w:val'), 'center')

    c_nil = rows[10].cells[8]
    c_nil.merge(rows[10].cells[9])
    c_nil.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(c_nil.paragraphs[0], headers[8], bold=True)
    v = c_nil._tc.get_or_add_tcPr().find(qn('w:vAlign'))
    if v is not None: v.set(qn('w:val'), 'center')

    # Row 11 (Item values)
    def c_txt(row_idx, idx, text, align="left"):
        c = rows[row_idx].cells[idx]
        if align == "right": c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if align == "center": c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(c.paragraphs[0], str(text))

    c_txt(11, 0, '-')
    c_txt(11, 1, _s(k.komoditi))
    c_txt(11, 2, _s(k.packaging))
    c_txt(11, 3, '-')
    c_txt(11, 4, _s(k.simbol))
    c_txt(11, 5, '-')
    c_txt(11, 6, f" {_id_fmt(k.volume, 2).rstrip('0').rstrip(',')} ", 'right')
    c_txt(11, 7, f" {_rp(k.harga_satuan).replace('Rp', '').strip()} ", 'right')
    c_txt(11, 8, ' Rp')
    c_txt(11, 9, f" {_rp(k.nilai_transaksi).replace('Rp', '').strip()} ", 'right')

    # Row 12 (PPN val)
    c_ppnl = rows[12].cells[0]
    c_ppnl.merge(rows[12].cells[7])
    c_ppnl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(c_ppnl.paragraphs[0], f"PPN {ppn_pct}% ", bold=True)
    _run(rows[12].cells[8].paragraphs[0], ' Rp')
    c_ppnr = rows[12].cells[9]
    c_ppnr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(c_ppnr.paragraphs[0], f" {_rp(k.nominal_ppn).replace('Rp', '').strip()} ", bold=True)

    # Row 13 (PPh val)
    c_pph22l = rows[13].cells[0]
    c_pph22l.merge(rows[13].cells[7])
    c_pph22l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pph_pct = invoice.pph_22_persen or 0
    _run(c_pph22l.paragraphs[0], f"PPh 22 ({pph_pct}%) ", bold=True)
    _run(rows[13].cells[8].paragraphs[0], ' Rp')
    c_pph22r = rows[13].cells[9]
    c_pph22r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pokok = (k.volume or 0) * (k.harga_satuan or 0) + (k.premi or 0)
    nom_pph = pokok * (pph_pct / 100)
    _run(c_pph22r.paragraphs[0], f" -{_id_fmt(nom_pph, 2).replace('Rp', '').strip()} ", bold=True)

    # Row 14 (Total)
    c_totl = rows[14].cells[0]
    c_totl.merge(rows[14].cells[7])
    c_totl.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(c_totl.paragraphs[0], "Jumlah Pembayaran ", bold=True)
    _run(rows[14].cells[8].paragraphs[0], ' Rp')
    c_totr = rows[14].cells[9]
    c_totr.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(c_totr.paragraphs[0], f" {_id_fmt(invoice.jumlah_pembayaran, 2).replace('Rp', '').strip()} ", bold=True)

    # Row 15 (Terbilang)
    c_terb = rows[15].cells[0]
    c_terb.merge(rows[15].cells[9])
    _run(c_terb.paragraphs[0], "Terbilang:\n", bold=True)
    _run(c_terb.paragraphs[0], f"{invoice.terbilang_invoice or _s(k.terbilang, '')} Rupiah", italic=True)

    # Row 16 (Transfer Ke)
    c_trans = rows[16].cells[0]
    c_trans.merge(rows[16].cells[9])
    _run(c_trans.paragraphs[0], "Transfer Ke:\n", bold=True)
    bank_info = _s(k.pembayaran_bank, 'Bank Rakyat Indonesia')
    rek_info = _s(k.pembayaran_rek_no, 'No. 0050-01-005356-30-0')
    atas_info = _s(k.pembayaran_atas_nama, 'PT Perkebunan Nusantara I Regional 8')
    _run(c_trans.paragraphs[0], f"{atas_info} {bank_info}\n")
    _run(c_trans.paragraphs[0], f"No. Rekening: {rek_info}")

    doc.add_paragraph()
    p_date = doc.add_paragraph(f"Makassar, {_date(invoice.tanggal_transaksi)}")
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_date.paragraph_format.space_before = Pt(24)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generate_do_docx(do) -> io.BytesIO:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin   = Cm(1.5)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    inv = do.invoice
    k   = inv.kontrak

    # Create ONE giant table for the entire Delivery Order
    tbl = doc.add_table(rows=7, cols=7)
    tbl.style = 'Table Grid'
    
    # 7 columns with precise widths totaling 18.0 cm
    CW7 = [3.5, 3.0, 2.5, 1.5, 3.5, 2.0, 2.0]
    
    # helper functions
    def _cw(cell, cm):
        tc = cell._tc
        pr = tc.get_or_add_tcPr()
        old = pr.find(qn('w:tcW'))
        if old is not None: pr.remove(old)
        w = OxmlElement('w:tcW')
        w.set(qn('w:w'), str(int(cm * 567)))
        w.set(qn('w:type'), 'dxa')
        pr.append(w)

    def _p0(cell, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = "" 
        p = cell.paragraphs[0]
        p.alignment = align
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        return p

    def _va(cell, val='center'):
        tc = cell._tc; pr = tc.get_or_add_tcPr()
        v = pr.find(qn('w:vAlign'))
        if v is None: v = OxmlElement('w:vAlign'); pr.append(v)
        v.set(qn('w:val'), val)

    # Force widths on all cells before merging
    for row in tbl.rows:
        for i, c in enumerate(row.cells):
            _cw(c, CW7[i])
            
    # --- MERGES ---
    tbl.cell(0,0).merge(tbl.cell(1,3))  # Row 0 & 1: Header (Company)
    tbl.cell(0,4).merge(tbl.cell(1,4))  # Row 0 & 1: Header (DELIVERY ORDER)
    tbl.cell(0,5).merge(tbl.cell(0,6))  # Row 0: No.
    tbl.cell(1,5).merge(tbl.cell(1,6))  # Row 1: Tanggal
    tbl.cell(2,0).merge(tbl.cell(2,6))  # Row 2: Kepada / Alamat
    tbl.cell(3,0).merge(tbl.cell(3,6))  # Row 3: INFO
    tbl.cell(6,0).merge(tbl.cell(6,6))  # Row 6: Catatan

    c_comp = tbl.cell(0,0)
    c_do   = tbl.cell(0,4)
    c_no   = tbl.cell(0,5)
    c_tgl  = tbl.cell(1,5)
    c_addr = tbl.cell(2,0)
    c_info = tbl.cell(3,0)
    c_cat  = tbl.cell(6,0)

    # 1. Company
    p_comp = _p0(c_comp, WD_ALIGN_PARAGRAPH.CENTER)
    _va(c_comp, 'center')
    _run(p_comp, 'PT PERKEBUNAN NUSANTARA I\n', bold=True, size=11)
    _run(p_comp, 'REGIONAL 8\n', bold=True, size=11)
    _run(p_comp, 'Jalan Urip Sumoharjo No 72-76, Kota Makassar', size=10)

    # 2. "DELIVERY ORDER"
    p_do = _p0(c_do, WD_ALIGN_PARAGRAPH.CENTER)
    _va(c_do, 'center')
    _run(p_do, 'DELIVERY ORDER', bold=True, size=11)

    # 3. No.
    p_no = _p0(c_no, WD_ALIGN_PARAGRAPH.LEFT)
    p_no.paragraph_format.left_indent = Pt(6)
    p_no.paragraph_format.space_before = Pt(3)
    _va(c_no, 'top')
    _run(p_no, 'No.\n', bold=True, size=10)
    _run(p_no, _s(do.no_do), size=10)

    # 4. Tanggal
    p_tgl = _p0(c_tgl, WD_ALIGN_PARAGRAPH.LEFT)
    p_tgl.paragraph_format.left_indent = Pt(6)
    p_tgl.paragraph_format.space_before = Pt(3)
    _va(c_tgl, 'top')
    _run(p_tgl, 'Tanggal\n', bold=True, size=10)
    d = do.tanggal_do
    tgl_str = f"{d.day}/{d.month}/{d.year}" if d else '-'
    _run(p_tgl, tgl_str, size=10)

    # 5. Kepada & Alamat Unit
    p_addr = _p0(c_addr, WD_ALIGN_PARAGRAPH.LEFT)
    p_addr.paragraph_format.left_indent = Pt(3)
    p_addr.paragraph_format.space_before = Pt(3)
    p_addr.paragraph_format.space_after = Pt(3)
    _va(c_addr, 'center')
    p_addr.paragraph_format.tab_stops.add_tab_stop(Cm(2.2))
    _run(p_addr, f"Kepada\t: Manajer Unit {_s(do.kepada_unit)}\n", size=10)
    _run(p_addr, f"Alamat\t: {_s(do.alamat_unit)}", size=10)

    # 6. Atas penyerahan...
    p_info = _p0(c_info, WD_ALIGN_PARAGRAPH.LEFT)
    p_info.paragraph_format.left_indent = Pt(3)
    p_info.paragraph_format.space_before = Pt(3)
    p_info.paragraph_format.space_after = Pt(3)
    _va(c_info, 'top')
    p_info.paragraph_format.tab_stops.add_tab_stop(Cm(2.2))
    p_info.paragraph_format.tab_stops.add_tab_stop(Cm(10.0))
    _run(p_info, 'Atas penyerahan D.O. harap diserahkan barang-barang tersebut di bawah ini:\n', size=10)
    _run(p_info, f"Kepada\t: {_s(k.pembeli)}\n", size=10)
    _run(p_info, f"Jenis Barang\t: {_s(k.komoditi)}\tTahun panen: {_s(k.tahun_panen)}", size=10)

    # 7. Header Row
    headers = ['Kebun', 'Jenis\nProduk/Mutu', 'Banyaknya\nBale/Karung', 'No.\nKav./\nChop', 'No. Kontrak', 'Berat Kotor/\nBersih', 'Levering']
    for i, h in enumerate(headers):
        c = tbl.cell(4, i)
        p = _p0(c, WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        _va(c, 'center')
        _run(p, h, bold=True, size=10)

    # 8. Data Row
    vol_str = _id_fmt(k.volume, 2).rstrip('0').rstrip(',') if (k.volume and k.volume > 0) else '-'
    bale_str = _id_fmt(k.banyaknya_bale_karung, 2).rstrip('0').rstrip(',') if (k.banyaknya_bale_karung and k.banyaknya_bale_karung > 0) else '-'

    data = [
        _s(k.kebun_produsen),
        _s(k.deskripsi_produk or k.komoditi),
        bale_str,
        _s(k.no_kav_chop, '-'),
        _s(k.no_kontrak),
        vol_str,
        _s(k.levering, '-')
    ]

    for i, val in enumerate(data):
        c = tbl.cell(5, i)
        _va(c, 'top')
        p = _p0(c, WD_ALIGN_PARAGRAPH.LEFT if i == 4 else WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(4)
        _run(p, val, size=10)
        if i == 0:
            _run(p, '\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n\n', size=10)

    # 9. Catatan
    p_cat = _p0(c_cat, WD_ALIGN_PARAGRAPH.LEFT)
    p_cat.paragraph_format.left_indent = Pt(3)
    p_cat.paragraph_format.space_before = Pt(3)
    _va(c_cat, 'top')
    _run(p_cat, 'CATATAN :\n', bold=True, size=10)
    _run(p_cat, _s(k.catatan, ''), size=10)
    _run(p_cat, '\n\n', size=10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
