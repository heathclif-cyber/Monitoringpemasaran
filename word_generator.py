import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

C = 'Calibri'
FS = 11  # base font size pt


# ── XML helpers ───────────────────────────────────────────────────────────────

def _no_borders(tbl_or_cell, is_table=True):
    sides = ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']
    if is_table:
        elem = tbl_or_cell._tbl
        pr_tag, bdr_tag = qn('w:tblPr'), qn('w:tblBorders')
        pr = elem.find(pr_tag)
        if pr is None:
            pr = OxmlElement('w:tblPr'); elem.insert(0, pr)
    else:
        elem = tbl_or_cell._tc
        pr = elem.get_or_add_tcPr()
        bdr_tag = qn('w:tcBdr')

    old = pr.find(bdr_tag)
    if old is not None:
        pr.remove(old)
    bdr = OxmlElement(bdr_tag[bdr_tag.index('}')+1:] if '}' in bdr_tag else bdr_tag)
    bdr = OxmlElement('w:tblBorders' if is_table else 'w:tcBdr')
    for s in sides:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), 'none'); b.set(qn('w:sz'), '0'); b.set(qn('w:color'), 'auto')
        bdr.append(b)
    pr.append(bdr)


def _cw(cell, cm):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    old = pr.find(qn('w:tcW'))
    if old is not None: pr.remove(old)
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(cm * 567)))
    w.set(qn('w:type'), 'dxa')
    pr.append(w)


def _valign_top(cell):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), 'top')
    pr.append(v)


# ── Text helpers ──────────────────────────────────────────────────────────────

def _run(para, text, bold=False, size=FS, underline=False, italic=False):
    r = para.add_run(text)
    r.bold = bold; r.underline = underline; r.italic = italic
    r.font.name = C; r.font.size = Pt(size)
    return r


def _para_fmt(p, sa=1):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(sa)


def _cp(cell):
    """Return cell's first paragraph, set compact spacing."""
    p = cell.paragraphs[0]
    _para_fmt(p)
    return p


# ── Format helpers ────────────────────────────────────────────────────────────

def _s(v, fb='-'):
    if v is None: return fb
    sv = str(v).strip()
    return sv if sv else fb


def _rp(v):
    if not v: return '-'
    return 'Rp' + '{:,.0f}'.format(v).replace(',', '.')


def _rp_full(v):
    if not v: return 'Rp0'
    return 'Rp' + '{:,.0f}'.format(v).replace(',', '.')


MONTHS = ['','Januari','Februari','Maret','April','Mei','Juni',
          'Juli','Agustus','September','Oktober','November','Desember']


def _date(d):
    if d is None: return '-'
    try: return f"{d.day} {MONTHS[d.month]} {d.year}"
    except: return str(d)


# ── Main generator ────────────────────────────────────────────────────────────

def generate_contract_docx(kontrak) -> io.BytesIO:
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2.5)
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2)

    sty = doc.styles['Normal']
    sty.font.name = C; sty.font.size = Pt(FS)

    k = kontrak

    # ── computed values ───────────────────────────────────────────────────────
    vol   = k.volume or 0
    harga = k.harga_satuan or 0
    premi = k.premi or 0
    ppn   = k.ppn_persen or 11
    sat   = _s(k.satuan, 'Unit')
    nilai = (vol * harga) + premi

    vol_str   = '{:,.0f}'.format(vol).replace(',', '.') + ' ' + sat if vol else '-'
    harga_str = _rp_full(harga) + ' per ' + sat if harga else '-'
    premi_str = _rp(premi)
    jml_str   = _rp_full(nilai) + (f' ({_s(k.terbilang)} Rupiah)' if k.terbilang else '')
    lama      = k.lama_pembayaran_hari or 15

    pjl_lines = _s(k.penjual).split('\n')
    pjl_nama  = pjl_lines[0].strip()
    pjl_addr  = pjl_lines[1].strip() if len(pjl_lines) > 1 else 'Jalan Urip Sumoharjo No. 72-76, Kota Makassar'

    pbl_nama = _s(k.pembeli, '[Nama Pembeli]')
    pbl_addr = _s(k.alamat_pembeli, '')
    if pbl_addr == '-': pbl_addr = ''

    syarat = [ln.strip() for ln in _s(k.syarat_syarat, '').split('\n') if ln.strip()]
    if not syarat:
        syarat = [
            'Pembayaran dilaksanakan selambat-lambatnya 15 hari kalender setelah ditandatanganinya Kontrak penjualan',
            'Penyerahan barang dilaksanakan setelah diterima bukti transfer pembayaran',
            'Pengambilan barang selambat-lambatnya 15 hari kalender dari batas akhir tanggal pembayaran',
            'Biaya-biaya yang terkait dengan administrasi Bank menjadi beban pembeli',
        ]

    dasar  = _s(k.dasar_ketentuan, 'Mengacu kepada tata Cara dan Ketentuan Penjualan Komoditi Perkebunan PT Perkebunan Nusantara III (Persero)')
    lokasi = _s(k.lokasi, 'Makassar')
    tgl    = _date(k.tanggal_kontrak)

    met    = _s(k.pembayaran_metode, 'Tunai')
    cara   = _s(k.pembayaran_cara, 'Transfer')
    bank   = _s(k.pembayaran_bank, 'Bank Rakyat Indonesia')
    atas   = _s(k.pembayaran_atas_nama, 'PT Perkebunan Nusantara I Regional 8')
    rek    = _s(k.pembayaran_rek_no, 'No. 0050-01-005356-30-0')

    # ── Title block ───────────────────────────────────────────────────────────
    def ctr(text, bold=False, ul=False, size=FS, sa=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_fmt(p, sa=sa)
        _run(p, text, bold=bold, underline=ul, size=size)

    ctr('KONTRAK PENJUALAN', bold=True, ul=True, size=12, sa=1)
    ctr(f'Nomor : {_s(k.no_kontrak)}', sa=1)
    ctr(f'Bid Offer Nomor : {_s(k.bid_offer_nomor, "")}', sa=6)

    # ── Main borderless table ─────────────────────────────────────────────────
    # 6 cols:  Label | : | Value | Label2 | : | Value2
    # widths:  3.8   |0.25| 4.8  | 2.8    |0.25| 3.6   = 15.5 cm
    CW = [3.8, 0.25, 4.8, 2.8, 0.25, 3.6]

    tbl = doc.add_table(rows=0, cols=6)
    _no_borders(tbl, is_table=True)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    def nr():
        """New row with borders removed and widths set."""
        row = tbl.add_row()
        for i, cell in enumerate(row.cells):
            _no_borders(cell, is_table=False)
            _cw(cell, CW[i])
            _valign_top(cell)
            _para_fmt(cell.paragraphs[0])
        return row

    def W(cell, text, bold=False):
        _run(_cp(cell), text, bold=bold)

    def COLON(cell):
        _run(_cp(cell), ':')

    # ── simple: label : value (merged cols 2-5) ───────────────────────────────
    def RS(label, value, bv=False):
        r = nr()
        W(r.cells[0], label, bold=True); COLON(r.cells[1])
        r.cells[2].merge(r.cells[5]); W(r.cells[2], value, bold=bv)

    # ── dual: label1:val1   label2:val2 ─────────────────────────────────────
    def RD(l1, v1, l2, v2, bv1=False, bv2=False):
        r = nr()
        W(r.cells[0], l1, bold=True); COLON(r.cells[1]); W(r.cells[2], v1, bold=bv1)
        W(r.cells[3], l2, bold=True); COLON(r.cells[4]); W(r.cells[5], v2, bold=bv2)

    # ── multiline value helper ────────────────────────────────────────────────
    def RM_multiline(label, lines_tuples):
        """lines_tuples = [(text, bold), ...]"""
        r = nr()
        W(r.cells[0], label, bold=True); COLON(r.cells[1])
        r.cells[2].merge(r.cells[5])
        p = _cp(r.cells[2])
        for i, (txt, b) in enumerate(lines_tuples):
            if i: p.add_run('\n')
            _run(p, txt, bold=b)

    # ── Pembayaran block helper ───────────────────────────────────────────────
    def pembayaran_block():
        # Row 1: Pembayaran : "Metode : X"   "Cara Pembayaran  : X"
        r1 = nr()
        W(r1.cells[0], 'Pembayaran', bold=True); COLON(r1.cells[1])
        p2 = _cp(r1.cells[2])
        _run(p2, 'Metode', bold=True); _run(p2, '    :    '); _run(p2, met)
        W(r1.cells[3], 'Cara\nPembayaran', bold=True); COLON(r1.cells[4]); W(r1.cells[5], cara)

        # Row 2: (empty)  "Nama Bank : X"   "Jatuh Tempo Pembayaran : Maksimal N Hari Kalender"
        r2 = nr()
        p2b = _cp(r2.cells[2])
        _run(p2b, 'Nama Bank', bold=True); _run(p2b, '   :   '); _run(p2b, bank)
        W(r2.cells[3], 'Jatuh Tempo\nPembayaran', bold=True); COLON(r2.cells[4])
        W(r2.cells[5], f'Maksimal {lama} Hari\nKalender')

        # Row 3: (empty)  "Atas Nama : X"
        r3 = nr()
        r3.cells[2].merge(r3.cells[5])
        p3 = _cp(r3.cells[2])
        _run(p3, 'Atas Nama', bold=True); _run(p3, '  :  '); _run(p3, atas)

        # Row 4: (empty)  "Rek No.  : X"
        r4 = nr()
        r4.cells[2].merge(r4.cells[5])
        p4 = _cp(r4.cells[2])
        _run(p4, 'Rek No.', bold=True); _run(p4, '    :    '); _run(p4, rek)

    # ── syarat multi-paragraph ────────────────────────────────────────────────
    def syarat_block():
        r = nr()
        W(r.cells[0], 'Syarat - Syarat Lain', bold=True); COLON(r.cells[1])
        r.cells[2].merge(r.cells[5])
        p = _cp(r.cells[2])
        alpha = 'abcdefgh'
        for i, s in enumerate(syarat):
            # strip any existing leading letter
            clean = s.strip()
            if len(clean) > 2 and clean[0].lower() in alpha and clean[1] in '.':
                clean = clean[2:].strip()
            letter = alpha[i] if i < len(alpha) else str(i+1)
            if i: p.add_run('\n')
            _run(p, f'{letter}.\t{clean}')

    # ── Populate rows ─────────────────────────────────────────────────────────
    RS('Pemilik Komoditas', _s(k.pemilik_komoditas), bv=True)
    RM_multiline('Penjual', [(pjl_nama, True), (pjl_addr, False)])
    lines_pb = [(pbl_nama, True)]
    if pbl_addr: lines_pb.append((pbl_addr, False))
    RM_multiline('Pembeli', lines_pb)
    RS('No. Referensi', _s(k.no_reff))
    RD('Komoditi', _s(k.komoditi), 'Jenis Komoditi', _s(k.jenis_komoditi))
    RD('Packaging', _s(k.packaging), 'Symbol', _s(k.simbol))
    RS('Deskripsi Produk', _s(k.deskripsi_produk))
    RS('Mutu', _s(k.mutu))
    RS('Produsen', _s(k.kebun_produsen))
    RS('Pelabuhan Muat', _s(k.pelabuhan_muat))
    RS('Volume', vol_str)
    RD('Harga Satuan', harga_str, 'Premi', premi_str)
    RS('PPN', f'Tarif Efektif {ppn}%')
    RS('Kondisi Penyerahan', _s(k.kondisi_penyerahan))
    pembayaran_block()
    RS('Waktu Penyerahan', _s(k.waktu_penyerahan))
    syarat_block()
    RS('Dasar Ketentuan', dasar)
    RS('Jumlah (Pokok)', jml_str)

    # ── Signature block ───────────────────────────────────────────────────────
    sp = doc.add_paragraph()
    _para_fmt(sp, sa=4)

    sig = doc.add_table(rows=4, cols=2)
    _no_borders(sig, is_table=True)
    sig.alignment = WD_TABLE_ALIGNMENT.LEFT

    for row in sig.rows:
        for cell in row.cells:
            _no_borders(cell, is_table=False)
            _para_fmt(cell.paragraphs[0])

    # Row 0: blank | Makassar, date (right aligned)
    p_date = sig.rows[0].cells[1].paragraphs[0]
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_date, f'{lokasi}, {tgl}')

    # Row 1: Persetujuan Pembeli | PT PTPN I Regional 8,
    pl = sig.rows[1].cells[0].paragraphs[0]
    _run(pl, 'Persetujuan Pembeli', bold=True)
    pr_ = sig.rows[1].cells[1].paragraphs[0]
    pr_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(pr_, 'PT Perkebunan Nusantara I Regional 8,', bold=True)

    # Row 2: signature space
    for cell in sig.rows[2].cells:
        _run(cell.paragraphs[0], '\n\n\n\n')

    # Row 3: name lines
    p_l3 = sig.rows[3].cells[0].paragraphs[0]
    _run(p_l3, '(________________________)')
    p_r3 = sig.rows[3].cells[1].paragraphs[0]
    p_r3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_r3, '(Pejabat Berwenang)')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
