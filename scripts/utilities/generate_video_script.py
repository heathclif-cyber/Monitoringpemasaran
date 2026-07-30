"""Generate promotional video script for Monitoring Pemasaran PTPN I."""
from __future__ import annotations

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUTPUT = r"D:\Script Aplikasi Pemasaran.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Arial"


def add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Arial"
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.name = "Arial"


def add_scene(
    doc: Document,
    *,
    no: str,
    waktu: str,
    visual: str,
    narasi: str,
    teks_layar: str = "",
    catatan: str = "",
) -> None:
    add_heading(doc, f"Adegan {no} — {waktu}", level=2)
    add_label_value(doc, "Visual / Aksi di layar", visual)
    add_label_value(doc, "Narasi (voice-over)", narasi)
    if teks_layar:
        add_label_value(doc, "Teks di layar (opsional)", teks_layar)
    if catatan:
        add_label_value(doc, "Catatan produksi", catatan)
    doc.add_paragraph()


def build() -> Document:
    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("SCRIPT VIDEO PROMOSI\nMonitoring Pemasaran PTPN I Regional 8")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "Arial"

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Durasi: 2 menit 30 detik  |  Format: Screen recording + voice-over")
    sr.font.size = Pt(11)
    sr.font.name = "Arial"
    sr.italic = True

    doc.add_paragraph()

    add_heading(doc, "Ringkasan Aplikasi", level=1)
    add_para(
        doc,
        "Monitoring Pemasaran adalah sistem otomasi dokumen penjualan dan pelaporan realisasi "
        "untuk PT Perkebunan Nusantara I Regional 8. Aplikasi ini menghasilkan Kontrak Penjualan, "
        "Invoice, Kuitansi, dan Delivery Order (DO) dalam format .docx baku perusahaan, "
        "menyediakan dashboard realisasi pemasaran, pencatatan persediaan, laporan digital terintegrasi, "
        "serta otomasi pengisian SPPn/SPPb di portal Superman — inovasi pertama di lingkungan PTPN "
        "yang menghubungkan data internal ke formulir eksternal secara otomatis (browser agent).",
    )

    add_heading(doc, "Spesifikasi Video", level=1)
    specs = [
        ("Judul kerja", "Monitoring Pemasaran — Pemasaran Digital PTPN I Regional 8"),
        ("Durasi target", "2:30 (boleh dipotong 2:00 atau diperpanjang 3:00)"),
        ("Audiens", "Manajemen Regional 8, tim pemasaran, keuangan, unit kebun"),
        ("Tone", "Profesional, modern, percaya diri — tonjolkan efisiensi & inovasi AI agent"),
        ("Musik", "Corporate uplifting, tempo sedang, fade di hook Superman"),
        ("Rekaman", "Screen capture aplikasi (Railway production atau lokal) + VO bahasa Indonesia"),
    ]
    for label, value in specs:
        add_label_value(doc, label, value)

    add_heading(doc, "Pesan Utama (3 pilar)", level=1)
    pillars = [
        "1. Satu sistem dari kontrak hingga DO — dokumen Word siap pakai, tanpa ketik ulang.",
        "2. Dashboard & Laporan Digital — realisasi pemasaran terpantau real-time per unit & komoditi.",
        "3. Otomasi Superman (AI Agent) — pengisian SPPn/SPPb otomatis dari data DO, pertama di PTPN.",
    ]
    for p in pillars:
        add_para(doc, p)

    add_heading(doc, "Alur Video (Timeline)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Waktu"
    hdr[1].text = "Segmen"
    hdr[2].text = "Menu / Fitur"
    rows = [
        ("0:00–0:20", "Hook + Dashboard", "Dashboard"),
        ("0:20–1:05", "Dokumen otomatis", "Buat Kontrak → Cetak Invoice → Delivery Order"),
        ("1:05–1:25", "Realisasi & stok", "Laporan Digital + Persediaan"),
        ("1:25–1:40", "Dukungan operasional", "Upload Dokumen · Berita Acara · Bypass"),
        ("1:40–2:20", "Hero: Superman AI Agent", "Buat SPPn Superman"),
        ("2:20–2:30", "Penutup", "Dashboard + tagline"),
    ]
    for waktu, segmen, fitur in rows:
        row = table.add_row().cells
        row[0].text = waktu
        row[1].text = segmen
        row[2].text = fitur

    doc.add_paragraph()
    add_heading(doc, "SCRIPT LENGKAP (Scene by Scene)", level=1)

    add_scene(
        doc,
        no="1",
        waktu="0:00 – 0:10",
        visual="Logo PTPN / judul aplikasi fade in. Cut ke sidebar menu lengkap: Dashboard, Buat Kontrak, "
        "Cetak Invoice, Delivery Order, Berita Acara, Upload Dokumen, Persediaan, Laporan Digital.",
        narasi="Di era digital perkebunan, tim pemasaran butuh lebih dari spreadsheet. "
        "Perkenalkan Monitoring Pemasaran PTPN I Regional 8 — satu platform untuk seluruh alur penjualan.",
        teks_layar="Monitoring Pemasaran | PTPN I Regional 8",
    )

    add_scene(
        doc,
        no="2",
        waktu="0:10 – 0:20",
        visual="Buka Dashboard. Tunjukkan kartu KPI: Pendapatan Omset, Cash In, Volume Realisasi. "
        "Scroll grafik Tren Bulanan dan Portofolio Komoditas. Filter Tahun / Unit / Komoditi.",
        narasi="Pantau realisasi pemasaran dalam satu layar. Omset, arus kas, volume terkirim, "
        "dan kelengkapan data SAP — terfilter per tahun, unit, dan komoditi.",
        teks_layar="Dashboard Realisasi Pemasaran",
        catatan="Highlight grafik garis & pie chart komoditas.",
    )

    add_scene(
        doc,
        no="3",
        waktu="0:20 – 0:40",
        visual="Menu Buat Kontrak. Wizard 3 langkah: Identitas → Komoditas → Finalisasi. "
        "Panel preview kanan menampilkan KONTRAK PENJUALAN live. Klik Simpan & Preview → Export .docx.",
        narasi="Buat kontrak penjualan dengan wizard tiga langkah. Volume, harga, PPN, PPh, "
        "dan terbilang terhitung otomatis. Preview langsung, export dokumen Word format baku perusahaan.",
        teks_layar="Buat Kontrak · Simpan & Preview · Export .docx",
        catatan="Opsional: tunjukkan multi-unit (beberapa kebun) atau tipe Payung + Berita Acara.",
    )

    add_scene(
        doc,
        no="4",
        waktu="0:40 – 0:55",
        visual="Menu Cetak Invoice. Pilih kontrak. Tunjukkan progress bar multi-invoice (% ter-invoice, Sisa). "
        "Isi Jumlah Pembayaran bertahap. Export .docx + Export Kuitansi.",
        narasi="Satu kontrak bisa punya banyak invoice untuk pembayaran bertahap. "
        "Sistem validasi total invoice dan menghasilkan Proforma Invoice plus Kuitansi proporsional — siap cetak.",
        teks_layar="Cetak Invoice · Export Kuitansi",
    )

    add_scene(
        doc,
        no="5",
        waktu="0:55 – 1:05",
        visual="Menu Delivery Order. Pilih invoice, isi Nominal Transfer. Volume DO dan Selisih ter-update otomatis. "
        "Tunjukkan cek Saldo Stok. Simpan → Export .docx.",
        narasi="Delivery Order terbit otomatis dari invoice. Volume DO proporsional, selisih pembayaran, "
        "dan cek persediaan — dihitung real-time sebelum dokumen DO di-export.",
        teks_layar="Delivery Order · Volume DO · Selisih",
    )

    add_scene(
        doc,
        no="6",
        waktu="1:05 – 1:18",
        visual="Menu Laporan Digital. Summary cards: Total Cash In, Pendapatan, Kekurangan Bayar, Sisa Barang. "
        "Scroll tabel — kolom No. DO, Invoice, Kontrak ter-freeze. Filter bulan berjalan. Export Excel.",
        narasi="Laporan Digital menyatukan kontrak, invoice, dan DO dalam satu tabel. "
        "Lacak sisa bayar, sisa volume, status SAP, dan deklarasi Superman — plus export Excel sekali klik.",
        teks_layar="Laporan Digital · Export Excel",
    )

    add_scene(
        doc,
        no="7",
        waktu="1:18 – 1:25",
        visual="Menu Persediaan. Input Persediaan Masuk per unit & jenis material. "
        "Tampilkan Saldo Saat Ini dan heatmap Peta. Klik Sinkron DO.",
        narasi="Menu Persediaan mencatat stok masuk per unit dan material. "
        "Saldo terintegrasi dengan DO — barang terkirim tercatat otomatis.",
        teks_layar="Persediaan · Saldo Saat Ini · Sinkron DO",
    )

    add_scene(
        doc,
        no="8",
        waktu="1:25 – 1:40",
        visual="Montage cepat: Upload Dokumen (kontrak/BA), Berita Acara payung, Input Bypass. "
        "Tunjukkan checklist dokumen hijau di Laporan.",
        narasi="Upload dokumen terpusat, Berita Acara untuk kontrak payung, "
        "dan Input Bypass untuk transaksi di luar kontrak — semua masuk dashboard dan laporan.",
        teks_layar="Upload Dokumen · Berita Acara · Input Bypass",
        catatan="B-roll 3–5 detik per menu, tanpa narasi panjang.",
    )

    add_scene(
        doc,
        no="9",
        waktu="1:40 – 1:55",
        visual="Kembali ke Laporan Digital. Pilih baris DO. Kolom Dok. Wajib hijau (Siap). "
        "Klik tombol SPPn Superman. Dialog konfirmasi muncul.",
        narasi="Dan inilah inovasi pertama di PTPN: otomasi Superman. "
        "Tanpa mengetik ulang di portal pajak, data DO Anda mengisi draft SPPn secara otomatis.",
        teks_layar="Status Deklarasi Superman · Buat SPPn Superman",
        catatan="Tonjolkan ini sebagai AI Agent / browser automation.",
    )

    add_scene(
        doc,
        no="10",
        waktu="1:55 – 2:10",
        visual="Dialog Login Superman (captcha). Progress bar: Membuka form → Mengisi baris SPPn → Menyimpan. "
        "Notifikasi sukses: SPPn + SPPb masuk To Do List. Kolom Superman terisi nomor.",
        narasi="Agen otomasi membaca kontrak, komoditi, nominal, dan dokumen pendukung — "
        "lalu mengisi form Superman: GL pendapatan per komoditi, PPh, BA, hingga simpan ke To Do List.",
        teks_layar="Membuat SPPn Superman… · Sudah SPPn",
        catatan="Jika demo live riskan, gunakan rekaman sukses sebelumnya (sawit/karet).",
    )

    add_scene(
        doc,
        no="11",
        waktu="2:10 – 2:20",
        visual="Split screen atau montage: Karet, Kelapa, Tebu, Sawit di Laporan. "
        "Tunjukkan mapping GL berbeda. Repository preview dokumen (ikon mata).",
        narasi="Mendukung semua komoditi Regional 8 — sawit, karet, kelapa, tebu, dan lainnya. "
        "Arsip kontrak, invoice, dan DO tersimpan dengan preview dokumen identik hasil download.",
        teks_layar="Multi-komoditi · Repository",
    )

    add_scene(
        doc,
        no="12",
        waktu="2:20 – 2:30",
        visual="Kembali ke Dashboard dengan data terbaru. Fade out ke logo + URL aplikasi Railway.",
        narasi="Dari kontrak sampai SPPn — satu sistem, tanpa ketik ulang. "
        "Monitoring Pemasaran PTPN I Regional 8. Pemasaran digital, cerdas, terintegrasi.",
        teks_layar="Monitoring Pemasaran · PTPN I Regional 8",
    )

    add_heading(doc, "Daftar Fitur untuk B-Roll (ceklis rekaman)", level=1)
    broll = [
        "Dashboard — filter tahun/unit/komoditi, grafik tren, Status Kelengkapan SAP",
        "Buat Kontrak — wizard 3 langkah, live preview, export .docx, multi-unit",
        "Cetak Invoice — multi-invoice, progress bar, export kuitansi",
        "Delivery Order — kalkulasi volume DO & selisih, cek stok, export .docx",
        "Berita Acara — kontrak payung, volume BA, bulan buku",
        "Upload Dokumen — kontrak, BA, deklarasi, kelengkapan 3/3",
        "Persediaan — input masuk, saldo, peta heatmap, sinkron DO",
        "Laporan Digital — summary KPI, tabel freeze column, edit SAP inline, export Excel",
        "Input Bypass — transaksi manual masuk laporan",
        "Repository — preview docx (ikon mata) kontrak/invoice/DO",
        "Superman — captcha dialog, progress 0–100%, notifikasi sukses, kolom Superman terisi",
    ]
    for item in broll:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "Tagline Alternatif (penutup)", level=1)
    taglines = [
        "Dari kontrak sampai SPPn — satu sistem, tanpa ketik ulang.",
        "Pemasaran digital PTPN I Regional 8 — dokumen otomatis, realisasi terpantau, Superman terisi sendiri.",
        "Inovasi pertama: AI Agent pengisian SPPn otomatis di lingkungan PTPN.",
        "Monitoring Pemasaran — efisien, akurat, terintegrasi.",
    ]
    for t in taglines:
        add_para(doc, f"• {t}")

    add_heading(doc, "Catatan Produksi", level=1)
    notes = [
        "Rekam dalam resolusi minimal 1920×1080, cursor smooth, hindari data sensitif (nominal bisa disamarkan).",
        "Pastikan minimal satu DO dengan dokumen pendukung ter-upload agar demo Superman berjalan.",
        "VO: kecepatan ±130 kata/menit. Total narasi ±320–350 kata untuk 2:30.",
        "Subtitle Indonesia disarankan untuk aksesibilitas di forum internal PTPN.",
        "End card: logo Regional 8 + nama aplikasi + kontak tim pengembang.",
    ]
    for n in notes:
        doc.add_paragraph(n, style="List Bullet")

    add_heading(doc, "Naskah Narasi Lengkap (continuous)", level=1)
    add_para(
        doc,
        "Di era digital perkebunan, tim pemasaran butuh lebih dari spreadsheet. "
        "Perkenalkan Monitoring Pemasaran PTPN I Regional 8 — satu platform untuk seluruh alur penjualan. "
        "Pantau realisasi pemasaran dalam satu layar: omset, arus kas, volume terkirim, dan kelengkapan data SAP. "
        "Buat kontrak penjualan dengan wizard tiga langkah — volume, harga, PPN, PPh, dan terbilang terhitung otomatis. "
        "Preview langsung, export dokumen Word format baku perusahaan. "
        "Satu kontrak bisa punya banyak invoice untuk pembayaran bertahap, plus kuitansi proporsional siap cetak. "
        "Delivery Order terbit otomatis dari invoice, dengan volume DO dan selisih pembayaran dihitung real-time. "
        "Laporan Digital menyatukan seluruh data, lacak sisa bayar dan status SAP, plus export Excel. "
        "Menu Persediaan mencatat stok per unit dan material, terintegrasi dengan DO. "
        "Upload dokumen, Berita Acara payung, dan Input Bypass melengkapi operasional harian. "
        "Dan inilah inovasi pertama di PTPN: otomasi Superman. "
        "Tanpa mengetik ulang di portal pajak, agen otomasi mengisi draft SPPn dari data DO Anda — "
        "mendukung semua komoditi Regional 8. "
        "Dari kontrak sampai SPPn — satu sistem, tanpa ketik ulang. "
        "Monitoring Pemasaran PTPN I Regional 8. Pemasaran digital, cerdas, terintegrasi.",
        italic=True,
    )

    # Footer
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Dibuat untuk promosi internal PTPN I Regional 8 · 2026")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    fr.font.name = "Arial"

    return doc


if __name__ == "__main__":
    doc = build()
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")